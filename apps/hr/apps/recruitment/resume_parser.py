"""AI-powered resume parser — extract candidate profile from PDF/DOC/DOCX/TXT.

Split into two deterministic stages so callers (and the Celery task) can cache the
extracted text and avoid re-reading the file:

  extract_text(bytes, name) -> str          # deterministic, no LLM
  parse_resume_text(text)   -> dict          # LLM structuring of already-extracted text
  parse_resume(bytes, name) -> dict          # convenience: extract then parse (legacy)
"""
from __future__ import annotations

import logging

logger = logging.getLogger(__name__)


def parse_resume(file_bytes: bytes, file_name: str = "") -> dict:
    """Extract + structure a resume in one call (kept for backward compatibility)."""
    text = extract_text(file_bytes, file_name)
    if not text:
        return {"error": "Could not extract text from file"}
    return parse_resume_text(text)


def parse_resume_text(text: str) -> dict:
    """Structure already-extracted resume text into candidate fields via the LLM.

    Returns a dict with: full_name, email, phone, current_company, current_role,
    total_experience_years, skills, education, experience, certifications, summary.
    Returns {"error": ...} on failure.
    """
    from django.conf import settings
    api_key = getattr(settings, "ANTHROPIC_API_KEY", "")
    if not api_key:
        return {"error": "ANTHROPIC_API_KEY not configured"}
    if not text:
        return {"error": "No resume text to parse"}

    try:
        import anthropic
        client = anthropic.Anthropic(api_key=api_key)
        model = getattr(settings, "ANTHROPIC_MODEL", "claude-sonnet-4-6")
        response = client.messages.create(
            model=model,
            max_tokens=1024,
            system=(
                "You are a resume data-extraction assistant for an HR recruitment system. "
                "Your ONLY task is to extract structured candidate data from the resume text provided "
                "and return it as valid JSON. "
                "Do NOT provide any commentary, hiring recommendations, or content outside the JSON response. "
                "Do NOT invent data not present in the resume — use empty string or null for missing fields. "
                "Do NOT evaluate or score the candidate."
            ),
            messages=[{
                "role": "user",
                "content": f"""Parse this resume and return ONLY valid JSON with these fields:
{{
  "full_name": "",
  "email": "",
  "phone": "",
  "current_company": "",
  "current_role": "",
  "total_experience_years": 0,
  "skills": [],
  "certifications": [],
  "education": [
    {{"degree": "", "institution": "", "year": ""}}
  ],
  "experience": [
    {{"company": "", "role": "", "from": "", "to": "", "description": ""}}
  ],
  "summary": ""
}}

Resume text:
{text[:6000]}"""
            }]
        )
        from .ai_utils import parse_llm_json
        result = parse_llm_json(response.content[0].text)
        return result or {"error": "Could not parse resume data"}
    except Exception as e:
        logger.exception("Resume parsing failed")
        return {"error": str(e)}


def extract_text(file_bytes: bytes, file_name: str) -> str:
    """Extract plain text from PDF / DOCX / DOC / TXT. Best-effort, never raises."""
    name_lower = (file_name or "").lower()

    if name_lower.endswith(".pdf"):
        try:
            import io

            import pdfplumber
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                return "\n".join(p.extract_text() or "" for p in pdf.pages)
        except ImportError:
            pass
        except Exception:
            return ""
        try:  # Fallback: pypdf (successor to PyPDF2)
            import io
            try:
                from pypdf import PdfReader
            except ImportError:
                from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(file_bytes))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception:
            return ""

    if name_lower.endswith(".docx"):
        try:
            import io

            import docx
            doc = docx.Document(io.BytesIO(file_bytes))
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception:
            return ""

    if name_lower.endswith(".doc"):
        # Legacy binary .doc — try python-docx (some .doc are actually docx), then
        # fall back to a lossy decode of the readable ASCII runs.
        try:
            import io

            import docx
            doc = docx.Document(io.BytesIO(file_bytes))
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception:
            try:
                return file_bytes.decode("latin-1", errors="ignore")
            except Exception:
                return ""

    # .txt and anything else: plain decode.
    try:
        return file_bytes.decode("utf-8", errors="ignore")
    except Exception:
        return ""


# Backwards-compatible alias (older callers import the private name).
_extract_text = extract_text
